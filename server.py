# Native libraries
from json import dumps
from os import getenv
from typing import Annotated, Literal, List, Tuple
from enum import Enum
from pathlib import Path
from io import BytesIO
import logging
logging.basicConfig(level=logging.INFO, force=True)
logger = logging.getLogger("GenFilesMCP")

# Third-party libraries
from pydantic import Field, BaseModel
from requests import post, get
from mcp.server.fastmcp import FastMCP, Context
from mcp.server.session import ServerSession
from docx import Document

# Utilities
from utils.load_md_templates import load_md_templates
from utils.upload_file import upload_file
from utils.download_file import download_file
from utils.knowledge import create_knowledge

# Parameters
URL = getenv('OWUI_URL')
if not URL:
    raise ValueError("OWUI_URL environment variable is required")
PORT = int(getenv('PORT', '8000'))
HWP_TEMPLATE, POWERPOINT_TEMPLATE, EXCEL_TEMPLATE, WORD_TEMPLATE,MARKDOWN_TEMPLATE, MCP_INSTRUCTIONS = load_md_templates()
# Enable or disable automatic creation of knowledge collections after upload
# Defaults to true to preserve existing behavior. Set to 'false' to disable.
ENABLE_CREATE_KNOWLEDGE = getenv('ENABLE_CREATE_KNOWLEDGE', 'true').lower() == 'true'

# Pydantic model for review comments
class ReviewComment(BaseModel):
    index: int
    comment: str

# Initialize FastMCP server
mcp = FastMCP(
    name = "GenFilesMCP",
    instructions = MCP_INSTRUCTIONS,   
    port = PORT,
    host = "0.0.0.0"
)

PRESENTON_ENDPOINT = getenv('PRESENTON_ENDPOINT')
if not PRESENTON_ENDPOINT:
    raise ValueError("PRESENTON_ENDPOINT environment variable is required")
PRESENTON_API_KEY = getenv('PRESENTON_API_KEY')
if not PRESENTON_API_KEY:
    raise ValueError("PRESENTON_API_KEY environment variable is required")
PRESENTON_BASE_URL = getenv('PRESENTON_BASE_URL')
if not PRESENTON_BASE_URL:
    raise ValueError("PRESENTON_BASE_URL environment variable is required")
## PPT 템플릿 우리껄로 수정
@mcp.tool(
    name = "generate_powerpoint",
    title = "Generate PowerPoint presentation",
    description = POWERPOINT_TEMPLATE
)
def generate_powerpoint(
    content: Annotated[
        str,
        Field(description="PPT 생성에 필요한 내용 요약")
    ],
    file_name: Annotated[
        str,
        Field(description="생성할 파일 이름 (확장자 제외).")
    ],
    user_id: Annotated[
        str,
        Field(description="Knowledge Base 등록용 유저 ID")
    ],
    template_type: Annotated[str, Field(description="PPT 템플릿 종류: general / modern / standard / swift", default="general")],
    ctx: Context[ServerSession, None]
) -> dict:

    """
    고도화 버전:
    - python_script는 더 이상 사용하지 않음
    - Presenton API로 PPT 생성
    - 기존 upload_file(), create_knowledge() 흐름은 그대로 유지
    """

    try:
        # [1] Presenton API 호출 (기본 템플릿 기반 PPT 생성)
        headers = {
            "Authorization": f"Bearer {PRESENTON_API_KEY}",
            "Content-Type": "application/json"
        }

        # python_script 대신 LLM이 제공할 수 있는 텍스트 기반 prompt 생성
        payload = {
            "content": content,   # backward compatibility 유지
            "n_slides": 8,
            "language": "ko",
            "template": template_type,  # 사용자가 선택한 템플릿 타입 사용
            "export_as": None
        }

        logger.info(f"Presenton API 호출: template={template_type}, slides={payload['n_slides']}")

        api_resp = post(PRESENTON_ENDPOINT, json=payload, headers=headers, timeout=600)
        
        # Log response for debugging
        if api_resp.status_code != 200:
            logger.error(f"API Error Response: {api_resp.text}")
        
        api_resp.raise_for_status()

        data = api_resp.json()
        logger.info(f"Presenton API 생성 응답: {data}")

        # Two-step process: generate then export
        presentation_id = data.get("presentation_id")
        if not presentation_id:
            error_msg = "Presentation ID not found in response"
            logger.error(f"[Presenton Error] {error_msg}")
            raise Exception(f"[Presenton Error] {error_msg}")

        logger.info(f"Presentation ID: {presentation_id}, calling export API...")

        # Call export API (bypasses Puppeteer)
        export_endpoint = PRESENTON_ENDPOINT.replace("/generate", "/export")
        export_payload = {
            "id": presentation_id,
            "export_as": None
        }

        export_resp = post(export_endpoint, json=export_payload, headers=headers, timeout=800)
        export_resp.raise_for_status()

        export_data = export_resp.json()
        logger.info(f"Export API 응답: {export_data}")

        file_path = export_data.get("path")
        file_url = export_data.get("file_url")

        if not file_path and not file_url:
            error_msg = export_data.get("error", "path 또는 file_url이 export 응답에 없습니다")
            logger.error(f"[Presenton Export Error] {error_msg}")
            raise Exception(f"[Presenton Export Error] {error_msg}")

        # [2] 생성된 PPT 파일 다운로드 또는 읽기
        buffer = BytesIO()
        
        if file_path and file_path.startswith("/app_data/"):
            # Local file path - read directly via Docker volume
            # presenton container's /app_data is mounted to host, accessible from gen_files_mcp
            import requests
            # Convert to presenton container's file endpoint
            file_download_url = f"{PRESENTON_BASE_URL}{file_path}"
            logger.info(f"PPT 파일 다운로드 시작 (via HTTP): {file_download_url}")
            file_resp = get(file_download_url, timeout=300)
            file_resp.raise_for_status()
            logger.info(f"PPT 파일 다운로드 완료: {len(file_resp.content)} bytes")
            buffer.write(file_resp.content)
        elif file_url:
            # Remote URL - download via HTTP
            logger.info(f"PPT 파일 다운로드 시작 (via URL): {file_url}")
            file_resp = get(file_url, timeout=300)
            file_resp.raise_for_status()
            logger.info(f"PPT 파일 다운로드 완료: {len(file_resp.content)} bytes")
            buffer.write(file_resp.content)
        else:
            raise Exception(f"Cannot download file from path: {file_path}")
        buffer.name = f"{file_name}.pptx"
        buffer.seek(0)

        # [4] 인증 헤더 가져오기 (기존 그대로)
        bearer_token = None
        try:
            bearer_token = ctx.request_context.request.headers.get("authorization")
        except:
            logger.error("Error retrieving authorization header")

        # [5] Open-WebUI 업로드 (기존 그대로)
        logger.info(f"Open-WebUI에 파일 업로드 시작: {file_name}.pptx")
        upload_result, request_data = upload_file(
            url=URL,
            token=bearer_token,
            file_data=buffer,
            filename=file_name,
            file_type="pptx"
        )

        if "error" in upload_result:
            logger.error(f"파일 업로드 실패: {upload_result['error']}")
            return upload_result

        logger.info(f"파일 업로드 성공: {upload_result.get('file_path_download', 'N/A')}")

        # [6] Knowledge Base 등록 (기존 그대로)
        if "file_path_download" in upload_result and ENABLE_CREATE_KNOWLEDGE:
            logger.info(f"Knowledge Base 등록 시작: user_id={user_id}")
            create_knowledge(
                url=URL,
                token=bearer_token,
                file_id=request_data["id"],
                user_id=user_id
            )
            logger.info("Knowledge Base 등록 완료")

        return upload_result

    except Exception as e:
        logger.error(f"PPT 생성 중 오류 발생: {str(e)}", exc_info=True)
        return dumps({
            "error": {
                "message": f"PPT 생성 실패: {str(e)}"
            }
        }, indent=4, ensure_ascii=False)

HWP_ENDPOINT = getenv('HWP_ENDPOINT')
if not HWP_ENDPOINT:
    raise ValueError("HWP_ENDPOINT environment variable is required")

@mcp.tool(
    name="generate_hwp",
    title="Generate HWP document",
    description=HWP_TEMPLATE
)
def generate_hwp(
    content: Annotated[str, Field(description="행정 문서 스타일의 HWP 문서 본문 텍스트 (제목/본문 포함)")],
    file_name: Annotated[str, Field(description="생성할 파일 이름 (확장자 제외)")],
    user_id: Annotated[str, Field(description="Knowledge Base 등록용 유저 ID")],
    template_type: Annotated[str, Field(description="HWP 템플릿 종류: default / v2", default="default")],
    ctx: Context[ServerSession, None]
) -> dict:

    try:
        # Authorization Header
        bearer_token = None
        try:
            bearer_token = ctx.request_context.request.headers.get("authorization")
        except:
            logger.error("Error retrieving authorization header")

        # Request Payload
        payload = {
            "text": content,
            "file_name": f"{file_name}.hwp",
            "template_type": template_type
        }

        logger.info(f"HWP API 호출: template_type={template_type}")

        api_resp = post(HWP_ENDPOINT, json=payload, timeout=600)
        api_resp.raise_for_status()

        # JSON인지 Binary인지 판단
        try:
            data = api_resp.json()
            logger.info(f"HWPX API JSON 응답: {data}")

            file_id = data.get("file_id")
            if not file_id:
                raise Exception("file_id missing in HWP API JSON response")

            download_url = f"{HWP_ENDPOINT.replace('/generate', '').rstrip('/')}/download/{file_id}"
            file_resp = get(download_url, timeout=600)
            file_resp.raise_for_status()
            file_bytes = file_resp.content

        except ValueError:
            logger.info("HWPX API returned binary file directly")
            file_bytes = api_resp.content

        # Buffer 생성
        buffer = BytesIO(file_bytes)
        buffer.name = f"{file_name}.hwp"
        buffer.seek(0)

        # Upload to Open-WebUI
        upload_result, request_data = upload_file(
            url=URL,
            token=bearer_token,
            file_data=buffer,
            filename=file_name,
            file_type="hwp"
        )

        if "error" in upload_result:
            logger.error(f"파일 업로드 실패: {upload_result['error']}")
            return upload_result

        logger.info(f"HWP 업로드 성공: {upload_result.get('file_path_download')}")

        # Knowledge 등록
        if "file_path_download" in upload_result and ENABLE_CREATE_KNOWLEDGE:
            create_knowledge(
                url=URL,
                token=bearer_token,
                file_id=request_data["id"],
                user_id=user_id
            )

        return upload_result

    except Exception as e:
        logger.error(f"HWP 생성 오류: {str(e)}", exc_info=True)
        return dumps({
            "error": {
                "message": f"HWP 생성 실패: {str(e)}"
            }
        }, indent=4, ensure_ascii=False)


@mcp.tool(
    name = "generate_excel",
    title = "Generate Excel workbook",
    description = EXCEL_TEMPLATE
)
async def generate_excel(
    python_script: Annotated[
        str, 
        Field(description="Complete Python script that generates the Excel workbook using the provided template.")
    ],
    file_name: Annotated[
        str, 
        Field(description="Desired name for the generated Excel file without the extension.")
    ],
    user_id: Annotated[
        str,
        Field(description="User ID to associate the knowledge base with the correct user.")
    ],
    ctx: Context[ServerSession, None]
) -> dict:
    """
    Generate an Excel file using a Python script.

    Returns:
        dict: Contains 'file_path_download' with a markdown hyperlink for downloading the generated Excel file.
              Format: "[Download {filename}.xlsx](/api/v1/files/{id}/content)"
    """
    try:
        # Create a buffer for the Excel file
        buffer = BytesIO()
        buffer.name = f'{file_name}.xlsx'
        context = {"xlsx_buffer": buffer}
        exec(python_script, context)

        # Reset buffer position to start
        buffer.seek(0)

        # Retrieve authorization header from the request context
        try:
            bearer_token = ctx.request_context.request.headers.get("authorization")
            logger.info(f"Recieved authorization header!")
        except:
            logger.error(f"Error retrieving authorization header")

        # Upload the generated Excel file
        response, request_data = upload_file(
            url=URL, 
            token=bearer_token, 
            file_data=buffer,
            filename=file_name,
            file_type="xlsx"
        )

        # If upload is successful, add to knowledge base
        if "file_path_download" in response and ENABLE_CREATE_KNOWLEDGE:
            # create knowledge base if not exists
            create_knowledge_status = create_knowledge(
                url=URL, 
                token=bearer_token,
                file_id=request_data['id'],
                user_id=user_id
            )
            if create_knowledge_status:
                logger.info("Knowledge base updated successfully.")
            else:
                logger.error(f"Error creating or updating knowledge base")
        elif "error" in response:
            logger.error(f"Error uploading the file.")
        else:
            logger.info("Skipping knowledge creation because ENABLE_CREATE_KNOWLEDGE is false")

        return response 
    
    except Exception as e:
        return dumps(
            {
                "error": {
                    "message": str(e)
                }
            }, 
            indent=4, 
            ensure_ascii=False
        )

@mcp.tool(
    name = "generate_word",
    title = "Generate Word document",
    description = WORD_TEMPLATE
)
async def generate_word(
    python_script: Annotated[
        str, 
        Field(description="Complete Python script that generates the Word document using the provided template.")
    ],
    file_name: Annotated[
        str, 
        Field(description="Desired name for the generated Word file without the extension.")
    ],
    user_id: Annotated[
        str,
        Field(description="User ID to associate the knowledge base with the correct user.")
    ],
    ctx: Context[ServerSession, None]
) -> dict:
    """
    Generate a Word file using a Python script.

    Returns:
        dict: Contains 'file_path_download' with a markdown hyperlink for downloading the generated Word file.
              Format: "[Download {filename}.docx](/api/v1/files/{id}/content)"
    """
    try:
        # Create a buffer for the Word file
        buffer = BytesIO()
        buffer.name = f'{file_name}.docx'
        context = {"docx_buffer": buffer}
        exec(python_script, context)

        # Reset buffer position to start
        buffer.seek(0)

        # Retrieve authorization header from the request context
        try:
            bearer_token = ctx.request_context.request.headers.get("authorization")
            logger.info(f"Recieved authorization header!")
        except:
            logger.error(f"Error retrieving authorization header")

        # Upload the generated Word file
        response, request_data = upload_file(
            url=URL, 
            token=bearer_token, 
            file_data=buffer,
            filename=file_name,
            file_type="docx"
        )

        # If upload is successful, add to knowledge base
        if "file_path_download" in response and ENABLE_CREATE_KNOWLEDGE:
            # create knowledge base if not exists
            create_knowledge_status = create_knowledge(
                url=URL, 
                token=bearer_token,
                file_id=request_data['id'],
                user_id=user_id
            )
            if create_knowledge_status:
                logger.info("Knowledge base updated successfully.")
            else:
                logger.error(f"Error creating or updating knowledge base")
        elif "error" in response:
            logger.error(f"Error uploading the file.")
        else:
            logger.info("Skipping knowledge creation because ENABLE_CREATE_KNOWLEDGE is false")

        return response 
    
    except Exception as e:
        return dumps(
            {
                "error": {
                    "message": str(e)
                }
            }, 
            indent=4, 
            ensure_ascii=False
        )

@mcp.tool(
    name = "generate_markdown",
    title = "Generate Markdown document",
    description = MARKDOWN_TEMPLATE
) 
async def generate_markdown(
    python_script: Annotated[
        str, 
        Field(description="Complete Python script that generates the Markdown document using the provided template.")
    ],
    file_name: Annotated[
        str, 
        Field(description="Desired name for the generated Markdown file without the extension.")
    ],
    user_id: Annotated[
        str,
        Field(description="User ID to associate the knowledge base with the correct user.")
    ],
    ctx: Context[ServerSession, None]
) -> dict:
    """
    Generate a Markdown file using a Python script.

    Returns:
        dict: Contains 'file_path_download' with a markdown hyperlink for downloading the generated Markdown file.
              Format: "[Download {filename}.md](/api/v1/files/{id}/content)"
    """
    try:
        # Create a buffer for the Markdown file
        buffer = BytesIO()
        buffer.name = f'{file_name}.md'
        context = {"md_buffer": buffer}
        exec(python_script, context)

        # Reset buffer position to start
        buffer.seek(0)

        # Retrieve authorization header from the request context
        try:
            bearer_token = ctx.request_context.request.headers.get("authorization")
            logger.info(f"Recieved authorization header!")
        except:
            logger.error(f"Error retrieving authorization header")

        # Upload the generated Markdown file
        response, request_data = upload_file(
            url=URL, 
            token=bearer_token, 
            file_data=buffer,
            filename=file_name,
            file_type="md"
        )

        # If upload is successful, add to knowledge base
        if "file_path_download" in response and ENABLE_CREATE_KNOWLEDGE:
            # create knowledge base if not exists
            create_knowledge_status = create_knowledge(
                url=URL, 
                token=bearer_token,
                file_id=request_data['id'],
                user_id=user_id
            )
            if create_knowledge_status:
                logger.info("Knowledge base updated successfully.")
            else:
                logger.error(f"Error creating or updating knowledge base")
        elif "error" in response:
            logger.error(f"Error uploading the file.")
        else:
            logger.info("Skipping knowledge creation because ENABLE_CREATE_KNOWLEDGE is false")

        return response 
    
    except Exception as e:
        return dumps(
            {
                "error": {
                    "message": str(e)
                }
            }, 
            indent=4, 
            ensure_ascii=False
        )
    
@mcp.tool(
    name="full_context_docx",
    title="Return the structure of a docx document",
    description="""Return the index, style and text of each element in a docx document. This includes paragraphs, headings, tables, images, and other components. The output is a JSON object that provides a detailed representation of the document's structure and content.
    The Agent will use this tool to understand the content and structure of the document before perform corrections (spelling, grammar, style suggestions, idea enhancements). Agent have to identify the index of each element to be able to add comments in the review_docx tool."""
)
async def full_context_docx(
    file_id: Annotated[
        str, 
        Field(description="ID of the existing docx file to analyze (from a previous chat upload).")
    ],
    file_name: Annotated[
        str, 
        Field(description="The name of the original docx file")
    ],
    ctx: Context[ServerSession, None]
) -> dict:
    """
    Return the structure of a docx document including index, style, and text of each element.
    Returns:
        dict: A JSON object with the structure of the document.
    """
    # Retrieve authorization header from the request context
    try:
        bearer_token = ctx.request_context.request.headers.get("authorization")
        logger.info(f"Recieved authorization header!")
    except:
        logger.error(f"Error retrieving authorization header")

    try:
        # Download in memory the docx file using the download_file helper
        docx_file = download_file(
            url=URL, 
            token=bearer_token, 
            file_id=file_id
        )

        if isinstance(docx_file, dict) and "error" in docx_file:
            return dumps(
                docx_file,
                indent=4,
                ensure_ascii=False
            )
        else:
            # Instantiate a Document object from the in-memory file
            doc = Document(docx_file)
            
            # Structure to return
            text_body = {
                "file_name": file_name,
                "file_id": file_id,
                "body": []
            }

            # list to store different parts of the document
            parts = []

            for idx, parts in enumerate(doc.paragraphs):
                # text of the paragraph
                text = parts.text.strip()

                if not text:
                    # skip empty paragraphs
                    continue  

                # style of the paragraph
                style = parts.style.name  
                text_body["body"].append({
                    "index": idx,
                    "style": style ,  # style.name
                    "text": text  # text
                })

            return dumps(
                text_body,
                indent=4,
                ensure_ascii=False
            )
    except Exception as e:
        return dumps(
            {
                "error": {
                    "message": str(e)
                }
            }, 
            indent=4, 
            ensure_ascii=False
        )

@mcp.tool(
    name="review_docx",
    title="Review and comment on docx document",
    description="""Review an existing docx document, perform corrections (spelling, grammar, style suggestions, idea enhancements), and add comments to cells. Returns a markdown hyperlink for downloading the reviewed file."""
)
async def review_docx(
    file_id: Annotated[
        str, 
        Field(description="ID of the existing docx file to review (from a previous chat upload).")
    ],
    file_name: Annotated[
        str, 
        Field(description="The name of the original docx file")
    ],
    review_comments: Annotated[
        List[ReviewComment], 
        Field(description="List of objects where each object has keys: 'index' (int) and 'comment' (str). Example: [{'index': 0, 'comment': 'Fix typo'}].")
    ],
    user_id: Annotated[
        str,
        Field(description="User ID to associate the knowledge base with the correct user.")
    ],
    ctx: Context[ServerSession, None]
) -> dict:
    """
    Review an existing docx document and add comments to specified elements.
    Returns:
        dict: Contains 'file_path_download' with a markdown hyperlink for downloading the reviewed docx file.
              Format: "[Download {filename}.docx](/api/v1/files/{id}/content)"
    """
    # Retrieve authorization header from the request context
    try:
        bearer_token = ctx.request_context.request.headers.get("authorization")
        logger.info(f"Recieved authorization header!")
    except:
        logger.error(f"Error retrieving authorization header")

    try:
        
        # Download the existing docx file
        docx_file = download_file(URL, bearer_token, file_id)
        if isinstance(docx_file, dict) and "error" in docx_file:
            return dumps(docx_file, indent=4, ensure_ascii=False)

        # Load the document
        doc = Document(docx_file)

        # Add comments to specified paragraphs
        paragraphs = list(doc.paragraphs)  # Get list of paragraphs
        for item in review_comments:
            try:
                index = item.index
                comment_text = item.comment
            except (AttributeError, TypeError):
                # malformed item; skip
                continue

            if index is None or comment_text is None:
                continue

            if 0 <= index < len(paragraphs):
                para = paragraphs[index]
                if para.runs:  # Ensure there are runs to comment on
                    # Add comment to the first run of the paragraph
                    doc.add_comment(
                        runs=[para.runs[0]],
                        text=comment_text,
                        author="AI Reviewer",
                        initials="AI"
                    )

        # Create a buffer for the reviewed file
        buffer = BytesIO()
        buffer.name = f'{Path(file_name).stem}_reviewed.docx'
        doc.save(buffer)
        buffer.seek(0)

        # Upload the reviewed docx file
        response, request_data = upload_file(
            url=URL, 
            token=bearer_token, 
            file_data=buffer,
            filename=f"{Path(file_name).stem}_reviewed",
            file_type="docx"
        )

        # If upload is successful, add to knowledge base
        if "file_path_download" in response and ENABLE_CREATE_KNOWLEDGE:
            # create knowledge base if not exists
            create_knowledge_status = create_knowledge(
                url=URL, 
                token=bearer_token,
                file_id=request_data['id'],
                user_id=user_id,
                knowledge_name="Documents Reviewed by AI"
            )
            if create_knowledge_status:
                logger.info("Knowledge base updated successfully.")
            else:
                logger.error(f"Error creating or updating knowledge base")
        elif "error" in response:
            logger.error(f"Error uploading the file.")
        else:
            logger.info("Skipping knowledge creation because ENABLE_CREATE_KNOWLEDGE is false")

        return response
    
    except Exception as e:
        return dumps(
            {
                "error": {
                    "message": str(e)
                }
            }, 
            indent=4, 
            ensure_ascii=False
        )
    
# Initialize and run the server
if __name__ == "__main__":
    mcp.run(
        transport="streamable-http"
    )


